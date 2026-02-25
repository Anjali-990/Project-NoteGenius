# myapp/views.py
import certifi
import ssl

import os
import io
import json
import random
import re
import tempfile
import logging
from urllib import request
from pydub import AudioSegment
AudioSegment.converter = "ffmpeg"
AudioSegment.ffprobe = "ffprobe"
from googletrans import Translator
translator = Translator()
from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, FileResponse, Http404
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q
from django.core.mail import send_mail
from django.utils import timezone
from transformers import pipeline
from rouge_score import rouge_scorer
last_generated_video_notes = ""
# models
from .models import Profile, EmailOTP, Notes

# file / office handling
import PyPDF2
from docx import Document as DocxReader
from docx import Document as DocxWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import fitz

import pytesseract


ssl._create_default_https_context = ssl.create_default_context
ssl._create_default_https_context = lambda: ssl.create_default_context(cafile=certifi.where())

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
MAX_WORDS = 20000
# optional audio libs
try:
    import speech_recognition as sr
    SR_AVAILABLE = True
except Exception:
    SR_AVAILABLE = False

# transformers summarizer / asr
import torch

logger = logging.getLogger(__name__)

summarizer = None
asr_pipeline = None
translator_pipeline = None

device = 0 if torch.cuda.is_available() else -1

qna_model = pipeline("text2text-generation", model="google/flan-t5-base", framework="pt", device=0)

try:
    summarizer = pipeline(
        "summarization",
        model="sshleifer/distilbart-cnn-12-6",
        device=device
    )
    print("✅ Summarizer loaded successfully")
except Exception as e:
    print("❌ Summarizer failed:", e)
    logger.exception("Summarizer initialization error")

try:
    asr_pipeline = pipeline(
        "automatic-speech-recognition",
        model="openai/whisper-base",
        device=device,
        torch_dtype=torch.float16 if device == 0 else torch.float32
    )
    print("✅ ASR loaded successfully")
except Exception as e:
    print("❌ ASR failed:", e)
    logger.exception("ASR initialization error")

try:
    translator_pipeline = pipeline(
        "translation",
        model="Helsinki-NLP/opus-mt-mul-en",
        device=device
    )
    print("✅ Translator loaded successfully")
except Exception as e:
    print("❌ Translator failed:", e)
    logger.exception("Translator initialization error")
# presets
LENGTH_PRESETS = {
    "short":  {"max_length": 120,  "min_length": 40},
    "medium": {"max_length": 250,  "min_length": 120},
    "long":   {"max_length": 450,  "min_length": 220},
}   

MAX_PROCESS_CHARS = 200000


# ---------------------------
# AUTH / PAGES
# ---------------------------


def home(request):
    return render(request, "home.html")


def login_view(request):
    if request.method == "POST":
        login_id = request.POST.get("login_id", "").strip()
        password = request.POST.get("password")
        otp_code = request.POST.get("login_otp", "").strip()

        user_obj = None
        if login_id:
            user_obj = User.objects.filter(Q(email__iexact=login_id) | Q(username__iexact=login_id)).first()

        # OTP login
        if otp_code and user_obj:
            otp_obj = EmailOTP.objects.filter(user=user_obj, code=otp_code, is_used=False).order_by("-created_at").first()
            if otp_obj and not otp_obj.is_expired():
                otp_obj.is_used = True
                otp_obj.save()
                login(request, user_obj)
                return redirect("home")
            else:
                return render(request, "login.html", {"error": "Invalid or expired OTP."})

        # password login
        if user_obj and password:
            user = authenticate(request, username=user_obj.username, password=password)
        else:
            user = None

        if user:
            login(request, user)
            return redirect("home")
        else:
            return render(request, "login.html", {"error": "Invalid email/username or password."})

    return render(request, "login.html")


def signup_view(request):
    if request.method == "POST":
        username = request.POST.get("username", "").strip()
        email = request.POST.get("email", "").strip()
        phone = request.POST.get("phone", "").strip()
        password1 = request.POST.get("password1", "")
        password2 = request.POST.get("password2", "")
        otp_input = (request.POST.get("signup_otp") or "").strip()

        # validations
        if password1 != password2:
            return render(request, "signup.html", {"error": "Passwords do not match."})
        if User.objects.filter(username=username).exists():
            return render(request, "signup.html", {"error": "Username already taken."})
        if User.objects.filter(email__iexact=email).exists():
            return render(request, "signup.html", {"error": "Email already registered."})

        # verify OTP created earlier for this email
        otp_obj = EmailOTP.objects.filter(email__iexact=email, code=otp_input, is_used=False).order_by("-created_at").first()
        if not otp_obj or otp_obj.is_expired():
            return render(request, "signup.html", {"error": "Invalid or expired OTP."})

        otp_obj.is_used = True
        otp_obj.save()

        # create user
        user = User.objects.create_user(username=username, email=email, password=password1)
        if phone:
            Profile.objects.create(user=user, phone=phone)

        messages.success(request, "Account created successfully! Please log in.")
        return redirect("login")

    return render(request, "signup.html")


def logout_view(request):
    logout(request)
    return redirect("login")


# ---------------------------
# OTP sender (works for login OR signup email-only)
# ---------------------------
@require_POST
def send_login_otp(request):
    login_id = (request.POST.get("login_id") or "").strip()
    if not login_id:
        return JsonResponse({"success": False, "message": "Please enter your email first."})

    # generate a fresh 6-digit code
    code = f"{random.randint(100000, 999999)}"

    # try find existing user by email or username (case-insensitive)
    user = User.objects.filter(Q(email__iexact=login_id) | Q(username__iexact=login_id)).first()

    # Invalidate any previous unused OTP rows for this user/email
    try:
        if user:
            EmailOTP.objects.filter(user=user, is_used=False).update(is_used=True)
        else:
            EmailOTP.objects.filter(email__iexact=login_id, is_used=False).update(is_used=True)
    except Exception:
        # non-fatal; log and continue
        logger.exception("Error invalidating previous OTPs")

    # create new OTP row; tie to user if present, otherwise store email only
    try:
        if user:
            otp_row = EmailOTP.objects.create(user=user, email=user.email, code=code)
            recipient = user.email
        else:
            otp_row = EmailOTP.objects.create(user=None, email=login_id, code=code)
            recipient = login_id
    except Exception:
        logger.exception("Failed to create OTP row")
        return JsonResponse({"success": False, "message": "Server error creating OTP."})

    # send mail
    subject = "Your NoteGenius OTP"
    message = f"Your OTP for NoteGenius is {code}. It will expire in 10 minutes."
    try:
        send_mail(subject, message, None, [recipient], fail_silently=False)
    except Exception as e:
        logger.exception("Failed to send OTP email")
        return JsonResponse({"success": False, "message": "Error sending OTP email. Check email settings."})

    return JsonResponse({"success": True, "message": f"OTP has been sent to {recipient}"})

# ---------------------------
#evaluate_summary
# ---------------------------
def evaluate_summary(original_text, generated_text):

    scorer = rouge_scorer.RougeScorer(['rouge1', 'rougeL'], use_stemmer=True)

    scores = scorer.score(original_text, generated_text)

    rouge1 = round(scores['rouge1'].fmeasure * 100, 2)
    rougeL = round(scores['rougeL'].fmeasure * 100, 2)

    original_words = len(original_text.split())
    summary_words = len(generated_text.split())

    compression_ratio = round((summary_words / original_words) * 100, 2) if original_words > 0 else 0

    return {
        "rouge1": rouge1,
        "rougeL": rougeL,
        "compression": compression_ratio
    }


# ---------------------------
# TEXT summarization helpers
# ---------------------------
def chunk_text_dynamic(text):
    length = len(text)
    if length < 5000:
        max_chars = 800
    elif length < 20000:
        max_chars = 1500
    else:
        max_chars = 2500
    chunks = []
    start = 0
    while start < length:
        end = min(start + max_chars, length)
        chunks.append(text[start:end])
        start = end
    return chunks


def generate_fast_notes(text, summary_length="medium"):

    if not summarizer:
        raise Exception("Summarizer model not loaded.")

    import re

    # 🔹 Clean junk
    text = re.sub(r'\b(\w+)( \1\b)+', r'\1', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    # 🔹 Hinglish detection
    hinglish_words = re.findall(
        r"\b(hai|tha|thi|kya|kaise|kyun|aur|lekin|matlab|krna|hona)\b",
        text.lower()
    )

    # 🔹 Translate to English if needed
    if translator_pipeline and len(hinglish_words) > 5:
        try:
            translated = translator_pipeline(text, max_length=3000)
            text = translated[0]["translation_text"]
        except:
            print("⚠ Translation failed")

    words = text.split()
    total_words = len(words)

    # 🔹 26% ratio
    max_chunk_words = 900
    chunks = []

    for i in range(0, total_words, max_chunk_words):
        chunk = " ".join(words[i:i + max_chunk_words])
        chunks.append(chunk)

    partial_summaries = []

    for chunk in chunks:
        chunk_word_count = len(chunk.split())
        chunk_target = int(chunk_word_count * 0.26)

        if chunk_target < 120:
            max_len = 150
        elif chunk_target < 300:
            max_len = 300
        elif chunk_target < 600:
            max_len = 600
        else:
            max_len = 800

        result = summarizer(
            chunk,
            max_length=max_len,
            min_length=int(max_len * 0.5),
            do_sample=False,
            truncation=True
        )

        partial_summaries.append(result[0]["summary_text"])

    final_text = " ".join(partial_summaries).strip()

    # 🔹 Sentence split
    sentences = re.split(r'(?<=[.!?]) +', final_text)
    sentences = [s.strip() for s in sentences if len(s.split()) > 5]

    # 🔹 Introduction (first 2 sentences)
    introduction = " ".join(sentences[:2])

    # 🔹 Conclusion (last 2 sentences)
    conclusion = " ".join(sentences[-2:])

    # 🔹 Important Bullet Points (long impactful sentences)
    important_bullets = [
        s for s in sentences
        if len(s.split()) > 18
    ][:8]

    # 🔹 Main Body (excluding intro + conclusion)
    body_sentences = sentences[2:-2]

    # Group body into logical sections
    sections = []
    temp = []

    for s in body_sentences:
        temp.append(s)
        if len(temp) >= 4:
            sections.append(" ".join(temp))
            temp = []

    if temp:
        sections.append(" ".join(temp))

    structured_output = ""

    # Introduction
    structured_output += "\n\nIntroduction\n\n" + introduction

    # Main Sections
    for i, section in enumerate(sections):
        structured_output += f"\n\nMain Topic {i+1}\n\n{section}"


    # Conclusion
    structured_output += "\n\nConclusion\n\n" + conclusion

    # 🔹 Key Terms (not counted in ratio)
    key_terms = list(dict.fromkeys(
        [w.strip(".,;:()").capitalize()
         for w in final_text.split()
         if len(w) > 6 and w.isalpha()]
    ))[:12]

    return {
        "short_summary": introduction,
        "bullets": important_bullets,
        "detailed_explanation": structured_output.strip(),
        "key_terms": key_terms
    }
# textup view
# ---------------------------
def textup(request):
    summary = {}
    error = ""
    warning = ""
    evaluation = {}
    input_text = ""

    if request.method == "POST":

        uploaded_file = request.FILES.get("file")
        pasted_text = request.POST.get("content", "").strip()

        # 1️⃣ FILE UPLOAD
        if uploaded_file:
            fname = uploaded_file.name.lower()

            if uploaded_file.size > 350 * 1024 * 1024:
                error = "File too large. Maximum allowed size is 350 MB."
            else:
                try:
                    if fname.endswith(".pdf"):
                        reader = PyPDF2.PdfReader(uploaded_file)
                        input_text = "\n".join(
                            [page.extract_text() or "" for page in reader.pages]
                        )

                    elif fname.endswith(".txt"):
                        input_text = uploaded_file.read().decode("utf-8", errors="ignore")

                    elif fname.endswith((".docx", ".doc")):
                        doc = DocxReader(uploaded_file)
                        input_text = "\n".join([p.text for p in doc.paragraphs])

                    else:
                        error = "Unsupported file type."

                except Exception as e:
                    logger.exception("File reading failed")
                    error = f"Error reading file: {str(e)}"

        # 2️⃣ PASTED TEXT
        elif pasted_text:
            input_text = pasted_text

        else:
            error = "Please paste some text or upload a supported file."

        # 🔥 LIMIT TO 20,000 WORDS
        if input_text:
            words = input_text.split()
            if len(words) > MAX_WORDS:
                input_text = " ".join(words[:MAX_WORDS])
                warning = "Text trimmed to first 20,000 words."

        # Optional translation
        if translator_pipeline and input_text and len(input_text) < 3000:
            try:
                translated = translator_pipeline(input_text, max_length=3000)
                input_text = translated[0]["translation_text"]
            except Exception:
                logger.exception("Translation failed")

        # 3️⃣ GENERATE SUMMARY
        if input_text.strip() and not error:
            try:
                summary_length = request.POST.get("summary_length", "medium")
                notes = generate_fast_notes(input_text, summary_length)

                request.session["latest_notes"] = notes
                request.session.modified = True

                summary = notes

                evaluation = evaluate_summary(
                    input_text,
                    summary["detailed_explanation"]
                )

            except Exception as e:
                logger.exception("Generation failed")
                error = str(e)

        last_generated_notes = summary.get("detailed_explanation", "")
        if "doubt" in request.POST:
            doubt_text = request.POST.get("doubt")

            if last_generated_notes and doubt_text:
                prompt = f"""
                Answer the question strictly using the notes below.

                Notes:
                {last_generated_notes}

                Question:
                {doubt_text}
                """

                result = qna_model(prompt, max_length=400, do_sample=False)

                return JsonResponse({
                    "doubt_answer": result[0]["generated_text"]
                })

            return JsonResponse({
                "doubt_answer": "No notes available."
            })

    return render(request, "textup.html", {
        "summary": summary,
        "error": error,
        "warning": warning,
        "evaluation": evaluation,
        "input_text": input_text
    })
# ---------------------------
# audio upload / ASR
# ---------------------------
@csrf_exempt
def audio_upload(request):
    if request.method == "GET":
        return render(request, "audioup.html")

    uploaded = None
    tmp_path = None

    try:
        if request.FILES.get("file"):
            uploaded = request.FILES["file"]
            suffix = os.path.splitext(uploaded.name)[1] or ".wav"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            for chunk in uploaded.chunks():
                tmp.write(chunk)
            tmp.flush()
            tmp.close()
            tmp_path = tmp.name

        elif request.FILES.get("recorded_blob"):
            uploaded = request.FILES["recorded_blob"]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
            for chunk in uploaded.chunks():
                tmp.write(chunk)
            tmp.flush()
            tmp.close()
            tmp_path = tmp.name
        else:
            return JsonResponse({"success": False, "error": "No audio received."})

        transcript = ""

        # ------------- ASR (Whisper) -------------
        if asr_pipeline is not None:
            try:
                out = asr_pipeline(tmp_path)
                if isinstance(out, dict) and "text" in out:
                    transcript = out["text"]
                elif isinstance(out, list) and out and "text" in out[0]:
                    transcript = out[0]["text"]
                else:
                    transcript = str(out)
            except Exception:
                logger.exception("ASR pipeline failed")
                transcript = ""

        # ------------- Google Fallback -------------
        if not transcript and SR_AVAILABLE:
            try:
                r = sr.Recognizer()
                with sr.AudioFile(tmp_path) as source:
                    audio = r.record(source)
                transcript = r.recognize_google(audio)
            except Exception:
                logger.exception("SpeechRecognition failed")
                transcript = ""

        if not transcript:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            return JsonResponse({
                "success": False,
                "error": "Transcription failed. Try WAV/MP3 or ensure ffmpeg is installed."
            })

        # ------------- Clean Transcript -------------
        if len(transcript) > MAX_PROCESS_CHARS:
            transcript = transcript[:MAX_PROCESS_CHARS]

        import re
        transcript = re.sub(r'\b(\w+)( \1\b)+', r'\1', transcript)
        transcript = re.sub(r'\b(ya|na|uh|um|hmm)\b', '', transcript, flags=re.IGNORECASE)
        transcript = re.sub(r'\s+', ' ', transcript).strip()

        # ------------- Generate Notes -------------
        summary_length = request.POST.get("summary_length", "medium")

        clean_input = f"""
        You are an academic note-making assistant.

        From the transcript below:
        - Create meaningful section headings based on actual topics.
        - DO NOT use generic titles like "Main Topic 1".
        - Use proper descriptive headings.
        - Provide structured student notes.

        Transcript:
        {transcript}
        """

        notes = generate_fast_notes(clean_input, summary_length)

        # ------------- ✅ EVALUATION ADDED -------------
        evaluation = evaluate_summary(
            transcript,
            notes["detailed_explanation"]
        )

        request.session["latest_notes"] = notes
        request.session.modified = True

        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

        last_generated_notes = notes["detailed_explanation"]
        if "doubt" in request.POST:
            doubt_text = request.POST.get("doubt")

            if last_generated_notes and doubt_text:
                prompt = f"""
        Answer the question strictly using the notes below.

        Notes:
        {last_generated_notes}

        Question:
        {doubt_text}
        """

            result = qna_model(prompt, max_length=400, do_sample=False)

            return JsonResponse({
                "doubt_answer": result[0]["generated_text"]
            })

        return JsonResponse({
            "doubt_answer": "No notes available."
        })

        return JsonResponse({
            "success": True,
            "summary": notes,
            "transcript": transcript,
            "evaluation": evaluation   # 🔥 added
        })

    except Exception:
        logger.exception("audio_upload error")
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        return JsonResponse({
            "success": False,
            "error": "Server error while processing audio."
        })

# ---------------------------
# video upload / ASR
# ---------------------------

@csrf_exempt
def video_upload(request):

    if request.method == "GET":
        return render(request, "videoup.html")

    # 🔥 1️⃣ HANDLE DOUBT FIRST
    if request.POST.get("doubt"):

        doubt_text = request.POST.get("doubt")
        last_notes = request.session.get("latest_notes")

        if not last_notes:
            return JsonResponse({
                "doubt_answer": "No video notes available."
            })

        prompt = f"""
        Answer strictly using the notes below.

        Notes:
        {last_notes.get("detailed_explanation", "")}

        Question:
        {doubt_text}
        """

        result = qna_model(prompt, max_length=400, do_sample=False)

        return JsonResponse({
            "doubt_answer": result[0]["generated_text"]
        })


    tmp_video_path = None
    tmp_audio_path = None

    try:
        # 🔥 2️⃣ HANDLE VIDEO UPLOAD
        if not request.FILES.get("file"):
            return JsonResponse({
                "success": False,
                "error": "No video received."
            })

        uploaded = request.FILES["file"]

        suffix = os.path.splitext(uploaded.name)[1] or ".mp4"
        tmp_video = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)

        for chunk in uploaded.chunks():
            tmp_video.write(chunk)

        tmp_video.close()
        tmp_video_path = tmp_video.name

        # 🔥 Extract audio
        from moviepy.editor import VideoFileClip

        clip = VideoFileClip(tmp_video_path)

        if clip.audio is None:
            clip.close()
            return JsonResponse({
                "success": False,
                "error": "No audio track found in video."
            })

        tmp_audio = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        tmp_audio_path = tmp_audio.name
        tmp_audio.close()

        clip.audio.write_audiofile(tmp_audio_path)
        clip.close()

        # 🔥 Transcribe
        transcript = ""

        if asr_pipeline:
            out = asr_pipeline(tmp_audio_path)
            transcript = out.get("text", "")

        if not transcript:
            return JsonResponse({
                "success": False,
                "error": "Transcription failed."
            })

        if len(transcript) > MAX_PROCESS_CHARS:
            transcript = transcript[:MAX_PROCESS_CHARS]

        # 🔥 Generate notes
        summary_length = request.POST.get("summary_length", "medium")
        notes = generate_fast_notes(transcript, summary_length)

        evaluation = evaluate_summary(
            transcript,
            notes["detailed_explanation"]
        )

        # 🔥 SAVE TO SESSION (very important)
        request.session["latest_notes"] = notes
        request.session.modified = True

        return JsonResponse({
            "success": True,
            "summary": notes,
            "transcript": transcript,
            "evaluation": evaluation
        })

    except Exception as e:
        logger.exception("Video processing failed")
        return JsonResponse({
            "success": False,
            "error": str(e)
        })

    finally:
        for path in [tmp_video_path, tmp_audio_path]:
            try:
                if path and os.path.exists(path):
                    os.unlink(path)
            except Exception:
                pass
#yt linkupload
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from youtube_transcript_api import YouTubeTranscriptApi
from django.shortcuts import render


@csrf_exempt
def yt_link_upload(request):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Only POST allowed."})

    video_url = request.POST.get("video_url")

    if not video_url:
        return JsonResponse({"success": False, "error": "No URL provided."})

    try:
        # 🔹 Extract video ID
        if "v=" in video_url:
            video_id = video_url.split("v=")[1].split("&")[0]
        elif "youtu.be/" in video_url:
            video_id = video_url.split("youtu.be/")[1].split("?")[0]
        else:
            return JsonResponse({"success": False, "error": "Invalid YouTube URL."})

        # 🔹 Get transcript (CORRECT WAY)
        from youtube_transcript_api import YouTubeTranscriptApi
        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)
        full_text = " ".join([item.text for item in transcript])

        if not full_text.strip():
            return JsonResponse({
                "success": False,
                "error": "Transcript not available for this video."
            })

        # 🔹 Limit words (same logic like text upload)
        words = full_text.split()
        MAX_WORDS = 20000
        if len(words) > MAX_WORDS:
            full_text = " ".join(words[:MAX_WORDS])

        # 🔹 Use your EXISTING summarizer function
        notes = generate_fast_notes(full_text, "medium")

        # 🔹 Save in session
        request.session["latest_notes"] = notes
        request.session.modified = True

        return JsonResponse({
            "success": True,
            "summary": notes
        })

    except Exception as e:
        return JsonResponse({
            "success": False,
            "error": str(e)
        })


def yt_link_page(request):
    return render(request, "ytlinkup.html")
#-------------------------------------

#-----------------
# downloads (read from session latest_notes)
# ---------------------------
def download_pdf(request):
    notes = request.session.get("latest_notes")
    if not notes:
        raise Http404("No generated notes available. Generate notes first.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    def add_section(title, content_list):
        header_style = styles["Heading2"]
        normal = styles["BodyText"]
        story.append(Paragraph(title, header_style))
        story.append(Spacer(1, 6))
        if isinstance(content_list, list):
            for item in content_list:
                story.append(Paragraph("• " + item, normal))
        else:
            story.append(Paragraph(content_list, normal))
        story.append(Spacer(1, 12))

    add_section("Short Summary", notes.get("short_summary", ""))
    add_section("Bullet Points", notes.get("bullets", []))
    add_section("Detailed Explanation", notes.get("detailed_explanation", ""))
    add_section("Key Terms", notes.get("key_terms", []))

    doc.build(story)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename="generated_notes.pdf")


def download_docx(request):
    notes = request.session.get("latest_notes")
    if not notes:
        raise Http404("No generated notes available. Generate notes first.")

    doc = DocxWriter()
    doc.add_heading("Short Summary", level=2)
    doc.add_paragraph(notes.get("short_summary", ""))
    doc.add_heading("Bullet Points", level=2)
    for b in notes.get("bullets", []):
        p = doc.add_paragraph()
        p.add_run("• " + b)
    doc.add_heading("Detailed Explanation", level=2)
    doc.add_paragraph(notes.get("detailed_explanation", ""))
    doc.add_heading("Key Terms", level=2)
    for t in notes.get("key_terms", []):
        p = doc.add_paragraph()
        p.add_run("• " + t)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename="generated_notes.docx")


# ---------------------------
# save notes (form or AJAX fetch)
# ---------------------------
@login_required
def save_notes(request):
    if request.method != "POST":
        return redirect("home")

    notes_data = request.POST.get("notes_data") or request.POST.get("notes_json")
    title = request.POST.get("title") or request.POST.get("notes_title") or "Untitled Note"
    note_type = request.POST.get("note_type") or request.POST.get("type") or "text"

    # If a JSON body was sent (fetch with JSON)
    if not notes_data and request.body:
        try:
            body = json.loads(request.body.decode("utf-8"))
            notes_data = body.get("notes_data") or body.get("notes_json")
            title = body.get("title", title)
            note_type = body.get("note_type", note_type)
        except Exception:
            # ignore parsing errors
            pass

    if not notes_data or not str(notes_data).strip():
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": False, "error": "No notes data provided."})
        messages.error(request, "Nothing to save.")
        return redirect(request.META.get("HTTP_REFERER", "home"))

    Notes.objects.create(user=request.user, title=title[:200], content=notes_data, note_type=note_type)

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse({"success": True})
    messages.success(request, "Notes saved successfully.")
    return redirect("savednotes")


# ---------------------------
# saved notes listing & CRUD (user-scoped)
# ---------------------------
@login_required
def savednotes(request):
    notes = Notes.objects.filter(user=request.user).order_by("-created_at")
    return render(request, "savednotes.html", {"notes": notes})


@login_required
def view_note(request, pk):
    note = get_object_or_404(Notes, pk=pk, user=request.user)
    return render(request, "view_note.html", {"note": note})


@login_required
def edit_note(request, pk):
    note = get_object_or_404(Notes, pk=pk, user=request.user)
    if request.method == "POST":
        title = request.POST.get("title", note.title)
        content = request.POST.get("content", note.content)
        note.title = title[:200]
        note.content = content
        note.updated_at = timezone.now()
        note.save()
        messages.success(request, "Note updated.")
        return redirect("view_note", pk=note.pk)
    return render(request, "edit_note.html", {"note": note})


@login_required
def delete_note(request, pk):
    note = get_object_or_404(Notes, pk=pk, user=request.user)
    if request.method == "POST":
        note.delete()
        messages.success(request, "Note deleted.")
        return redirect("savednotes")
    return render(request, "confirm_delete.html", {"note": note})


# ---------------------------
# QnA page
# ---------------------------

from transformers import pipeline
from django.http import JsonResponse
from django.shortcuts import render
from .models import Notes   # 👈 make sure your saved notes model name is correct

# ---------------------------
# Helper: Extract Text
# ---------------------------

def extract_text(file):
    if file.name.endswith(".pdf"):
        text = ""
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        for page in pdf:
            text += page.get_text()
        return text

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")

    return ""


# ---------------------------
# Helper: Chunking (Reduced to 1500 for model memory limits)
# ---------------------------

def split_into_chunks(text, chunk_size=1500):
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]


# ---------------------------
# Context storage (for doubts)
# ---------------------------

stored_context = ""
last_generated_qna = ""


# ---------------------------
# Main View
# ---------------------------

def qna_upload(request):
    global stored_context, last_generated_qna

    if request.method == "GET":
        return render(request, "QnA.html")

    if request.method == "POST":

        # --------------------
        # 1️⃣ SAVE TO NOTES (Untouched)
        # --------------------
        if "save_qna" in request.POST:
            if last_generated_qna and request.user.is_authenticated:
                Notes.objects.create(
                    user=request.user,
                    title="QnA Notes",
                    content=last_generated_qna
                )
                return JsonResponse({"status": "saved"})
            return JsonResponse({"status": "error"})


        # --------------------
        # 2️⃣ DOUBT CLEAR (Untouched as requested)
        # --------------------
        if "doubt" in request.POST:
            doubt_text = request.POST.get("doubt")

            if stored_context and doubt_text:
                prompt = f"""
Answer strictly using ONLY the original document below.

If the answer is not present, say:
"The answer is not available in the provided document."

Original Document:
{stored_context}

Question:
{doubt_text}

Provide a clear explanation.
"""

                result = qna_model(prompt, max_length=700, do_sample=False)

                return JsonResponse({
                    "doubt_answer": result[0]["generated_text"]
                })

            return JsonResponse({
                "doubt_answer": "No document context available."
            })

        # --------------------
        # 3️⃣ QnA GENERATION (T5-Optimized)
        # --------------------
        uploaded_file = request.FILES.get("file")
        text_input = request.POST.get("text")
        content = extract_text(uploaded_file) if uploaded_file else text_input

        if content:
            stored_context = content
            # We use smaller chunks (1000 chars) to ensure the model focuses
            chunks = split_into_chunks(content, chunk_size=1000)
            
            all_qna = ""
            count = 1

            for chunk in chunks:

                if count > 10:
                    break

                # 🔥 Clean citations & junk before sending to model
                clean_chunk = re.sub(r'\(.*?\)', '', chunk)   # remove brackets
                clean_chunk = re.sub(r'http\S+', '', clean_chunk)  # remove links
                clean_chunk = re.sub(r'\s+', ' ', clean_chunk).strip()

                prompt = f"""
Generate ONE clear academic question and its detailed answer 
based strictly on the text below.

Format strictly as:

Question:
<question>

Answer:
<answer>

Text:
{clean_chunk}
"""

                result = qna_model(
                    prompt,
                    max_length=300,
                    min_length=80,
                    do_sample=False,        # 🔥 deterministic = less garbage
                    repetition_penalty=1.3
                )

                generated = result[0]["generated_text"]

                # 🔥 Force structure
                if "Question:" not in generated:
                    continue
                if "Answer:" not in generated:
                    continue

                all_qna += f"Q{count}. {generated}\n\n---\n\n"
                count += 1

            last_generated_qna = all_qna
            evaluation = evaluate_qna(all_qna, stored_context)

            return JsonResponse({
                "success": True,
                "qna": all_qna,
                "evaluation": evaluation
            })

        return JsonResponse({"qna": "No valid content provided."})

import re

def evaluate_qna(qna_text, original_text):
    questions = re.findall(r"Q[:\d]*\s*(.*?)[\n]", qna_text)
    answers = re.findall(r"Answer:\s*(.*?)(?=\nQ|$)", qna_text, re.DOTALL)

    total_q = len(questions)
    unique_q = len(set(questions))

    # redundancy score
    redundancy = round((unique_q / total_q) * 100, 2) if total_q else 0

    # simple consistency check
    match_count = 0
    for ans in answers:
        if ans.strip()[:30] in original_text:
            match_count += 1

    consistency = round((match_count / len(answers)) * 100, 2) if answers else 0

    # coverage proxy
    coverage = round((len(answers) / total_q) * 100, 2) if total_q else 0

    return {
        "redundancy": redundancy,
        "consistency": consistency,
        "coverage": coverage
    }

@csrf_exempt
def quiz_upload(request):

    if request.method == "GET":
        return render(request, "Quiz.html")

    if request.method == "POST":

        uploaded_file = request.FILES.get("file")
        text_input = request.POST.get("text")
        content = ""

        # 1️⃣ TEXT INPUT
        if text_input:
            content = text_input

        # 2️⃣ FILE HANDLING
        elif uploaded_file:

            name = uploaded_file.name.lower()

            # PDF
            if name.endswith(".pdf"):
                content = extract_text(uploaded_file)

            # DOCX / TXT
            elif name.endswith((".docx", ".doc", ".txt")):
                content = extract_text(uploaded_file)

            # IMAGE → OCR
            elif name.endswith((".png", ".jpg", ".jpeg")):
                image = Image.open(uploaded_file)
                content = pytesseract.image_to_string(image)

            # AUDIO
            elif name.endswith((".mp3", ".wav", ".m4a")):
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                for chunk in uploaded_file.chunks():
                    tmp.write(chunk)
                tmp.close()
                result = asr_pipeline(tmp.name)
                content = result.get("text", "")

            # VIDEO
            elif name.endswith((".mp4", ".mkv", ".mov")):
                from moviepy.editor import VideoFileClip

                tmp_video = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
                for chunk in uploaded_file.chunks():
                    tmp_video.write(chunk)
                tmp_video.close()

                clip = VideoFileClip(tmp_video.name)

                tmp_audio = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                clip.audio.write_audiofile(tmp_audio.name)
                clip.close()

                result = asr_pipeline(tmp_audio.name)
                content = result.get("text", "")

        if not content:
            return JsonResponse({"success": False, "error": "No content extracted."})

        # 3️⃣ GENERATE QUIZ
        prompt = f"""
Generate 10 multiple choice questions from the text below.

Return STRICT JSON format:
[
 {{
  "question": "...",
  "options": ["A","B","C","D"],
  "correct": 0
 }}
]

Text:
{content[:3000]}
"""

        result = qna_model(prompt, max_length=1500, do_sample=False)

        try:
            quiz_json = json.loads(result[0]["generated_text"])
        except:
            return JsonResponse({"success": False, "error": "Quiz generation failed."})

        return JsonResponse({
            "success": True,
            "quiz": quiz_json
        })


def profile(request):
    return render(request, "profile.html")


def about(request):
    return render(request, "aboutus.html")


def feedback(request):
    return render(request, "feedback.html")


def settings_page(request):
    return render(request, "settings.html")

from django.contrib.auth.decorators import login_required
from django.utils import timezone
from .models import Profile, Notes


@login_required
def profile(request):
    user = request.user

    # get or create profile
    profile, created = Profile.objects.get_or_create(user=user)

    if request.method == "POST":
        profile.occupation = request.POST.get("occupation", "")
        profile.usage = request.POST.get("usage", "")
        profile.bio = request.POST.get("bio", "")
        profile.save()

        messages.success(request, "Profile updated successfully!")
        return redirect("profile")

    context = {
        "profile": profile,
        "total_notes": Notes.objects.filter(user=user).count(),
        "total_quizzes": 0,  # jab QuizRecord add hoga tab connect karenge
        "total_qna": 0,
    }
    return render(request, "profile.html", context)